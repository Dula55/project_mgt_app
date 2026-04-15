import os
import gc
import re
import socket
import secrets
import hashlib
import threading
from io import BytesIO
from pathlib import Path
from functools import wraps
from datetime import datetime, timedelta

from dotenv import load_dotenv
from flask import (
    Flask, render_template, request, redirect, url_for, send_file, flash,
    session, jsonify, send_from_directory, abort
)
from flask_mail import Mail, Message
from flask_migrate import Migrate
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from models import db, User

try:
    from apscheduler.schedulers.background import BackgroundScheduler
except Exception:
    BackgroundScheduler = None

# ================= ENV LOADING =================
env_path = Path(__file__).parent / ".env"
if env_path.exists():
    load_dotenv(dotenv_path=env_path, override=True)
else:
    load_dotenv(override=True)

# ================= PROCESS SETTINGS =================
os.environ["MPLBACKEND"] = "Agg"
os.environ["MALLOC_ARENA_MAX"] = "2"
os.environ["OMP_NUM_THREADS"] = "1"
threading.stack_size(1024 * 512)

# ================= APP INIT =================
app = Flask(__name__)

# To share the same data across devices, use one shared database.
# Set DATABASE_URL to PostgreSQL/MySQL on your server for best results.
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL", "sqlite:///projects.db")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = os.environ.get("UPLOAD_FOLDER", "static/uploads")
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", secrets.token_hex(32))
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SECURE"] = True  # only if using HTTPS
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"

def str_to_bool(value):
    return str(value).strip().lower() in ("true", "1", "yes")

# ================= MAIL CONFIG =================
MAIL_USERNAME = os.getenv("MAIL_USERNAME")
MAIL_PASSWORD = os.getenv("MAIL_PASSWORD")
MAIL_SERVER = os.getenv("MAIL_SERVER", "smtp.gmail.com")
MAIL_PORT = os.getenv("MAIL_PORT", "587")
MAIL_USE_TLS = os.getenv("MAIL_USE_TLS", "True")
MAIL_USE_SSL = os.getenv("MAIL_USE_SSL", "False")

app.config["MAIL_USERNAME"] = MAIL_USERNAME
app.config["MAIL_PASSWORD"] = MAIL_PASSWORD
app.config["MAIL_SERVER"] = MAIL_SERVER
try:
    app.config["MAIL_PORT"] = int(MAIL_PORT)
except ValueError:
    app.config["MAIL_PORT"] = 587
app.config["MAIL_USE_TLS"] = str_to_bool(MAIL_USE_TLS)
app.config["MAIL_USE_SSL"] = str_to_bool(MAIL_USE_SSL)
app.config["MAIL_DEFAULT_SENDER"] = MAIL_USERNAME or "noreply@projectmanagement.com"

mail = Mail(app)

# ================= MODELS =================
from models import db, Project, Task, MediaFile, TeamMember, project_team

db.init_app(app)
migrate = Migrate(app, db)

scheduler = BackgroundScheduler() if BackgroundScheduler else None

# ================= TASK TEMPLATES =================
ROAD_PHASES = ["Site Preparation", "Earth Work", "Drainage Construction", "Asphalt Laying"]

BUILDING_PHASES = [
    "Site Preparation",
    "Setting Out & Foundation",
    "Block & Form Work",
    "Roofing",
    "Plastering & Painting",
]

COMMERCIAL_DEVELOPMENT_TASKS = [
    {"name": "Mobilization to site", "activity": "Site preparation", "dependencies": "Site possession", "category": "Preliminary Works", "duration": 3},
    {"name": "Demolition of existing structure", "activity": "Existing building removal", "dependencies": "Mobilization to site", "category": "Demolition", "duration": 5},
    {"name": "Setting out", "activity": "Layout positioning", "dependencies": "Demolition of existing structure", "category": "Site Works", "duration": 2},
    {"name": "Excavation & soil testing", "activity": "Foundation excavation", "dependencies": "Setting out", "category": "Earthworks", "duration": 7},
    {"name": "Blinding", "activity": "Lean concrete", "dependencies": "Excavation & soil testing", "category": "Foundation", "duration": 2},
    {"name": "Footing reinforcement", "activity": "Reinforcement installation", "dependencies": "Blinding", "category": "Foundation", "duration": 3},
    {"name": "Footing casting (Pad)", "activity": "Concrete placement", "dependencies": "Footing reinforcement", "category": "Foundation", "duration": 2},
    {"name": "Beam & beam casting", "activity": "Structural tie beams", "dependencies": "Footing casting (Pad)", "category": "Structural Frame", "duration": 4},
    {"name": "Column formwork", "activity": "Formwork works", "dependencies": "Beam & beam casting", "category": "Structural Frame", "duration": 3},
    {"name": "Column reinforcement", "activity": "Steel fixing", "dependencies": "Column formwork", "category": "Structural Frame", "duration": 4},
    {"name": "Column casting", "activity": "Concrete works", "dependencies": "Column reinforcement", "category": "Structural Frame", "duration": 3},
]

ALLOWED_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "gif", "webm", "webp"}
ALLOWED_VIDEO_EXTENSIONS = {"mp4", "mov", "avi", "mkv", "webm"}

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# ================= HELPERS =================
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please login first.', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def parse_date(value):
    if not value:
        return None
    if isinstance(value, datetime):
        return value.date()
    return datetime.strptime(value, "%Y-%m-%d").date()

def allowed_file(filename, file_type="image"):
    ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""
    if file_type == "image":
        return ext in ALLOWED_IMAGE_EXTENSIONS
    return ext in ALLOWED_VIDEO_EXTENSIONS

def calculate_end_date(start_date, duration_days):
    if start_date and duration_days is not None:
        return start_date + timedelta(days=int(duration_days))
    return None

def get_project_duration_days(project):
    if project and project.start_date and project.end_date:
        return max(0, (project.end_date - project.start_date).days)
    return 0

def task_cost(task):
    planned = float(task.planned_cost or 0)
    if planned <= 0:
        planned = float((task.duration_days or 0) * 1000)
    actual = float(task.actual_cost or 0)
    if actual <= 0:
        actual = planned * (float(task.progress or 0) / 100.0)
    return planned, actual

def safe_commit():
    try:
        db.session.commit()
        return True
    except Exception:
        db.session.rollback()
        raise

def project_report_payload(project):
    tasks = Task.query.filter_by(project_id=project.id).all()
    members = project.team_members

    task_rows = []
    for t in tasks:
        planned, actual = task_cost(t)
        task_rows.append({
            "id": t.id,
            "name": t.name,
            "progress": float(t.progress or 0),
            "start_date": t.start_date.strftime("%Y-%m-%d") if t.start_date else None,
            "end_date": t.end_date.strftime("%Y-%m-%d") if t.end_date else None,
            "duration_days": t.duration_days or 0,
            "planned_cost": planned,
            "actual_cost": actual,
            "task_category": t.task_category or "Uncategorized",
            "assigned_to": t.assigned_to.name if t.assigned_to else "Unassigned",
            "assigned_to_id": t.assigned_to_id,
        })

    member_rows = []
    for m in members:
        assigned = [t for t in tasks if t.assigned_to_id == m.id]
        total = len(assigned)
        completed = len([t for t in assigned if float(t.progress or 0) >= 100])
        delayed = len([t for t in assigned if t.end_date and project.end_date and t.end_date > project.end_date])
        avg_progress = round(sum(float(t.progress or 0) for t in assigned) / total, 1) if total else 0
        member_rows.append({
            "id": m.id,
            "name": m.name,
            "email": m.email,
            "total_tasks": total,
            "completed_tasks": completed,
            "delayed_tasks": delayed,
            "avg_progress": avg_progress,
        })

    return task_rows, member_rows

def send_invitation_email(email, project_name, project_id, invite_token):
    if not app.config["MAIL_USERNAME"] or not app.config["MAIL_PASSWORD"]:
        return False
    try:
        with app.app_context():
            invite_link = url_for("accept_invitation", token=invite_token, _external=True)
            sender_email = app.config["MAIL_DEFAULT_SENDER"]
            msg = Message(
                subject=f"Invitation to join project: {project_name}",
                sender=sender_email,
                recipients=[email],
                reply_to=sender_email,
            )
            msg.body = f"""You have been invited to join the project: {project_name}

Click the link below to accept the invitation:
{invite_link}

This invitation will expire in 7 days.
"""
            msg.html = f"""
            <html>
              <body style="font-family: Arial, sans-serif;">
                <h2>You've been invited!</h2>
                <p>You have been invited to join the project: <strong>{project_name}</strong></p>
                <p><a href="{invite_link}">Accept Invitation</a></p>
                <p>This invitation expires in 7 days.</p>
              </body>
            </html>
            """
            socket.setdefaulttimeout(30)
            mail.send(msg)
            return True
    except Exception as e:
        print(f"Invitation email failed for {email}: {e}")
        return False

def generate_pdf_report(project):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    tasks = Task.query.filter_by(project_id=project.id).all()
    story.append(Paragraph(f"{project.name} Weekly Report", styles["Title"]))
    story.append(Paragraph(f"Overall Completion: {project.completion:.1f}%", styles["Normal"]))
    story.append(Spacer(1, 12))

    table_data = [["Task", "Progress", "Planned Cost", "Actual Cost"]]
    for t in tasks:
        planned, actual = task_cost(t)
        table_data.append([t.name, f"{float(t.progress or 0):.1f}%", f"{planned:.2f}", f"{actual:.2f}"])

    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))
    story.append(table)
    doc.build(story)
    buffer.seek(0)
    return buffer

def send_weekly_reports():
    projects = Project.query.all()
    for project in projects:
        recipients = [m.email for m in project.team_members if m.email]
        if not recipients:
            continue
        try:
            pdf = generate_pdf_report(project)
            msg = Message(
                subject=f"Weekly Report - {project.name}",
                recipients=recipients,
                body="Attached is the weekly project report."
            )
            msg.attach(f"{project.name}_weekly.pdf", "application/pdf", pdf.read())
            mail.send(msg)
        except Exception as e:
            print(f"Weekly report failed for {project.name}: {e}")

# ================= BASIC ROUTES =================
@app.route("/health")
def health_check():
    return "OK", 200

@app.route("/test_email")
def test_email():
    if not app.config["MAIL_USERNAME"] or not app.config["MAIL_PASSWORD"]:
        return jsonify({"status": "error", "message": "Email not configured"}), 500
    try:
        msg = Message(
            subject="Test Email from Project Management System",
            sender=app.config["MAIL_DEFAULT_SENDER"],
            recipients=[app.config["MAIL_USERNAME"]],
            body="This is a test email to verify email configuration."
        )
        mail.send(msg)
        return jsonify({"status": "success", "message": "Test email sent"})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/")
def index():
    try:
        member_id = session.get("member_id")
        if member_id:
            member = TeamMember.query.get(member_id)
            # Show projects the member is part of
            projects = member.projects if member else []
        else:
            # When not logged in, show all projects (read-only mode)
            projects = Project.query.all()

        for project in projects:
            for task in project.tasks:
                task.media_files = MediaFile.query.filter_by(task_id=task.id).all()

        return render_template("index.html", projects=projects, session=session)
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Index error: {e}")
        return render_template("index.html", projects=[], session=session), 200

# ================= API (GOOD FOR MULTI-DEVICE ACCESS) =================
@app.route("/api/projects")
def api_projects():
    member_id = session.get("member_id")
    if member_id:
        member = TeamMember.query.get(member_id)
        projects = member.projects if member else []
    else:
        projects = Project.query.all()
    
    return jsonify({
        "success": True,
        "projects": [
            {
                "id": p.id,
                "name": p.name,
                "project_type": p.project_type,
                "start_date": p.start_date.strftime("%Y-%m-%d") if p.start_date else None,
                "end_date": p.end_date.strftime("%Y-%m-%d") if p.end_date else None,
                "completion": round(p.completion, 1),
                "tasks_count": len(p.tasks),
            }
            for p in projects
        ]
    })

@app.route("/api/projects", methods=["POST"])
def api_create_project():
    data = request.get_json(force=True, silent=True) or request.form
    try:
        name = data.get("name")
        project_type = data.get("type") or data.get("project_type")
        if not name or not project_type:
            return jsonify({"success": False, "message": "name and project_type are required"}), 400

        project = Project(
            name=name,
            project_type=project_type,
            start_date=parse_date(data.get("start_date")),
            end_date=parse_date(data.get("end_date")),
        )
        db.session.add(project)
        safe_commit()
        
        # If user is logged in, automatically add them as a team member
        member_id = session.get("member_id")
        if member_id:
            member = TeamMember.query.get(member_id)
            if member and member not in project.team_members:
                project.team_members.append(member)
                safe_commit()

        return jsonify({"success": True, "message": "Project created successfully", "project_id": project.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500

@app.route("/api/projects/<int:project_id>/tasks")
def api_project_tasks(project_id):
    project = Project.query.get_or_404(project_id)
    tasks = Task.query.filter_by(project_id=project.id).order_by(Task.created_at.asc()).all()
    return jsonify({
        "success": True,
        "project": {
            "id": project.id,
            "name": project.name,
            "project_type": project.project_type,
            "completion": round(project.completion, 1),
        },
        "tasks": [
            {
                "id": t.id,
                "name": t.name,
                "progress": float(t.progress or 0),
                "start_date": t.start_date.strftime("%Y-%m-%d") if t.start_date else None,
                "end_date": t.end_date.strftime("%Y-%m-%d") if t.end_date else None,
                "duration_days": t.duration_days,
                "planned_cost": float(t.planned_cost or 0),
                "actual_cost": float(t.actual_cost or 0),
                "assigned_to": t.assigned_to.name if t.assigned_to else None,
                "task_category": t.task_category,
            }
            for t in tasks
        ]
    })

@app.route("/api/projects/<int:project_id>/tasks", methods=["POST"])
def api_create_task(project_id):
    project = Project.query.get_or_404(project_id)
    data = request.get_json(force=True, silent=True) or request.form
    try:
        name = data.get("name")
        if not name:
            return jsonify({"success": False, "message": "Task name is required"}), 400

        task = Task(
            name=name,
            project_id=project.id,
            progress=float(data.get("progress", 0) or 0),
            activity_description=data.get("activity_description"),
            dependencies=data.get("dependencies"),
            task_category=data.get("task_category"),
            start_date=parse_date(data.get("start_date")),
            end_date=parse_date(data.get("end_date")),
            duration_days=int(data["duration_days"]) if data.get("duration_days") not in (None, "") else None,
            planned_cost=float(data.get("planned_cost", 0) or 0),
            actual_cost=float(data.get("actual_cost", 0) or 0),
        )

        if data.get("assigned_to_id"):
            task.assigned_to_id = int(data.get("assigned_to_id"))

        if task.start_date and task.duration_days and not task.end_date:
            task.end_date = calculate_end_date(task.start_date, task.duration_days)

        db.session.add(task)
        safe_commit()
        return jsonify({"success": True, "message": "Task created successfully", "task_id": task.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500

# ================= PROJECT / TASK UI ROUTES =================
@app.route("/add_project", methods=["POST"])
def add_project():
    try:
        name = request.form["name"]
        project_type = request.form["type"]
        start_date_obj = parse_date(request.form.get("start_date"))
        end_date_obj = parse_date(request.form.get("end_date"))

        project = Project(name=name, project_type=project_type, start_date=start_date_obj, end_date=end_date_obj)
        db.session.add(project)
        db.session.flush()
        
        # If user is logged in, automatically add them as a team member
        member_id = session.get("member_id")
        if member_id:
            member = TeamMember.query.get(member_id)
            if member and member not in project.team_members:
                project.team_members.append(member)

        if project_type == "Road":
            phases, detailed = ROAD_PHASES, False
        elif project_type == "Building":
            phases, detailed = BUILDING_PHASES, False
        elif project_type == "Commercial":
            phases, detailed = COMMERCIAL_DEVELOPMENT_TASKS, True
        else:
            phases, detailed = [], False

        total_days = (end_date_obj - start_date_obj).days if start_date_obj and end_date_obj and end_date_obj > start_date_obj else 200
        current_start = start_date_obj

        if not phases:
            flash("Project created successfully, but no template tasks were added.", "warning")
        elif not detailed:
            phase_duration = max(1, total_days // len(phases))
            for phase in phases:
                task_end = calculate_end_date(current_start, phase_duration) if current_start else None
                db.session.add(Task(
                    name=phase,
                    progress=0.0,
                    project_id=project.id,
                    start_date=current_start,
                    end_date=task_end,
                    duration_days=phase_duration if current_start else None,
                    planned_cost=float(phase_duration * 1000),
                    actual_cost=0.0,
                ))
                current_start = task_end
        else:
            for i, task_data in enumerate(phases):
                task_duration = int(task_data.get("duration", 3))
                task_end = calculate_end_date(current_start, task_duration) if current_start else None
                db.session.add(Task(
                    name=task_data["name"],
                    activity_description=task_data.get("activity"),
                    dependencies=task_data.get("dependencies"),
                    task_category=task_data.get("category"),
                    progress=0.0,
                    project_id=project.id,
                    start_date=current_start,
                    end_date=task_end,
                    duration_days=task_duration if current_start else None,
                    planned_cost=float(task_duration * 1000 * (1 + i * 0.03)),
                    actual_cost=0.0,
                ))
                current_start = task_end

        safe_commit()
        flash(f'Project "{name}" created successfully!', "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error creating project: {str(e)}", "error")

    return redirect(url_for("index"))

@app.route("/edit_project/<int:project_id>", methods=["POST"])
def edit_project(project_id):
    project = Project.query.get_or_404(project_id)
    try:
        project.name = request.form.get("name", project.name)
        project.project_type = request.form.get("type", project.project_type)
        if request.form.get("start_date"):
            project.start_date = parse_date(request.form.get("start_date"))
        if request.form.get("end_date"):
            project.end_date = parse_date(request.form.get("end_date"))
        safe_commit()
        flash("Project updated successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error updating project: {str(e)}", "error")
    return redirect(url_for("index"))

@app.route("/report/<int:project_id>")
def report_view(project_id):
    project = Project.query.get_or_404(project_id)
    tasks, members = project_report_payload(project)
    return render_template("report.html", project=project, tasks=tasks, members=members)

@app.route("/project_members/<int:project_id>")
def project_members(project_id):
    project = Project.query.get_or_404(project_id)
    return render_template("project_members.html", project=project)

@app.route("/invite_members/<int:project_id>", methods=["POST"])
def invite_members(project_id):
    project = Project.query.get_or_404(project_id)
    emails_input = request.form.get("emails", "")
    is_ajax = request.headers.get("X-Requested-With") == "XMLHttpRequest"

    emails = re.split(r"[,\\n;]+", emails_input)
    emails = [email.strip().lower() for email in emails if email.strip()]

    if not emails:
        msg = "Please enter at least one email address"
        if is_ajax:
            return jsonify({"success": False, "message": msg}), 400
        flash(msg, "error")
        return redirect(request.referrer or url_for("project_members", project_id=project.id))

    invited_count, failed, success_emails = 0, [], []

    for email in emails:
        if not re.match(r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$", email):
            failed.append(f"{email} (invalid format)")
            continue

        try:
            member = TeamMember.query.filter_by(email=email).first()
            if not member:
                member = TeamMember(email=email, name=email.split("@")[0])
                db.session.add(member)
                db.session.flush()

            if member in project.team_members:
                failed.append(f"{email} (already a member)")
                continue

            token_data = f"{member.id}_{project.id}_{datetime.utcnow().timestamp()}_{secrets.token_hex(8)}"
            invite_token = hashlib.sha256(token_data.encode()).hexdigest()
            member.invite_token = invite_token
            member.token_expiry = datetime.utcnow() + timedelta(days=7)
            project.team_members.append(member)
            safe_commit()

            if send_invitation_email(email, project.name, project.id, invite_token):
                invited_count += 1
                success_emails.append(email)
            else:
                project.team_members.remove(member)
                member.invite_token = None
                member.token_expiry = None
                safe_commit()
                failed.append(f"{email} (email sending failed)")
        except Exception as e:
            db.session.rollback()
            failed.append(f"{email} ({str(e)})")

    if is_ajax:
        return jsonify({
            "success": invited_count > 0,
            "invited_count": invited_count,
            "failed_count": len(failed),
            "success_emails": success_emails,
            "failed_emails": failed,
        })

    if invited_count > 0:
        flash(f"Successfully invited {invited_count} member(s).", "success")
    if failed:
        flash(f"Some invitations failed: {', '.join(failed[:5])}", "warning")
    return redirect(request.referrer or url_for("project_members", project_id=project.id))

@app.route("/accept_invitation/<token>")
def accept_invitation(token):
    member = TeamMember.query.filter_by(invite_token=token).first()
    if not member:
        flash("Invalid or expired invitation link.", "error")
        return redirect(url_for("index"))

    if member.token_expiry and datetime.utcnow() > member.token_expiry:
        flash("Invitation link has expired.", "error")
        return redirect(url_for("index"))

    session["member_id"] = member.id
    session["member_name"] = member.name
    session["member_email"] = member.email
    session.permanent = True

    member.invite_token = None
    member.token_expiry = None
    safe_commit()

    flash(f"Welcome {member.name or member.email}!", "success")
    return redirect(url_for("index"))

# ================= REGISTER/LOGIN/LOGOUT ROUTES =================
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':

        email = (request.form.get('email') or '').strip().lower()
        name = (request.form.get('name') or '').strip()
        password = request.form.get('password')
        role = request.form.get('role', 'team_member')

        # validation
        if not email or not password or not name:
            flash("Email, name and password are required.", "error")
            return redirect(url_for('register'))

        existing = User.query.filter_by(email=email).first()
        if existing:
            flash("Email already exists. Please login.", "error")
            return redirect(url_for('login'))

        try:
            user = User(
                email=email,
                name=name,
                password_hash=generate_password_hash(password),
            )

            # store role if your model supports it
            if hasattr(User, "role"):
                user.role = role

            db.session.add(user)
            db.session.commit()

            # ✅ AUTO LOGIN
            session.clear()
            session['user_id'] = user.id
            session['user_email'] = user.email
            session['user_role'] = role

            flash("Account created and logged in successfully!", "success")

            # redirect based on role
            if role == "admin":
                return redirect(url_for("admin"))
            else:
                return redirect(url_for("index"))

        except Exception as e:
            db.session.rollback()
            flash(f"Registration failed: {str(e)}", "error")
            return redirect(url_for('register'))

    return render_template("register.html")


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email').strip().lower()
        password = request.form.get('password')

        # Validate input
        if not email or not password:
            flash('Please enter both email and password.', 'error')
            return redirect(url_for('login'))

        user = User.query.filter_by(email=email).first()

        # Check credentials
        if user and check_password_hash(user.password_hash, password):
            session['user_id'] = user.id
            session['user_email'] = user.email

            flash('Login successful!', 'success')
            return redirect(url_for('index'))  # or dashboard

        else:
            flash('Invalid email or password.', 'error')
            return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))


@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')


@app.route("/admin", methods=["GET", "POST"])
def admin():
    """Admin portal for managing team members"""
    # Check if user is logged in and is admin
    if not session.get("member_id") or not session.get("is_admin"):
        flash("Access denied. Admin privileges required.", "error")
        return redirect(url_for("index"))
    
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        name = request.form.get("name", "").strip()
        
        if not email or not name:
            flash("Both email and name are required", "error")
            return redirect(url_for("admin"))
        
        # Validate email format
        if not re.match(r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$", email):
            flash("Please enter a valid email address", "error")
            return redirect(url_for("admin"))
        
        # Check if user already exists
        existing_member = TeamMember.query.filter_by(email=email).first()
        if existing_member:
            flash(f"User with email {email} already exists", "error")
            return redirect(url_for("admin"))
        
        # Create new team member (not admin by default)
        new_member = TeamMember(email=email, name=name)
        db.session.add(new_member)
        safe_commit()
        
        flash(f"Team member {name} ({email}) has been registered successfully!", "success")
        return redirect(url_for("admin"))
    
    # GET request - show admin dashboard
    members = TeamMember.query.all()
    total_projects = Project.query.count()
    total_tasks = Task.query.count()
    admin_emails = os.environ.get("ADMIN_EMAILS", "").split(",")
    admin_emails = [email.strip().lower() for email in admin_emails if email.strip()]
    
    # Prepare member list with admin status
    member_list = []
    for member in members:
        member_list.append({
            "id": member.id,
            "name": member.name,
            "email": member.email,
            "joined_at": member.joined_at,
            "is_admin": member.email in admin_emails
        })
    
    return render_template("admin.html", 
                         members=member_list,
                         total_members=len(members),
                         total_projects=total_projects,
                         total_tasks=total_tasks,
                         admin_count=len([m for m in members if m.email in admin_emails]))

@app.route("/admin/delete_member/<int:member_id>", methods=["POST"])
def delete_member(member_id):
    """Delete a team member (admin only)"""
    # Check if user is logged in and is admin
    if not session.get("member_id") or not session.get("is_admin"):
        flash("Access denied. Admin privileges required.", "error")
        return redirect(url_for("index"))
    
    # Don't allow deleting yourself
    if member_id == session.get("member_id"):
        flash("You cannot delete your own account.", "error")
        return redirect(url_for("admin"))
    
    member = TeamMember.query.get_or_404(member_id)
    
    # Check if member is admin (prevent deleting other admins)
    admin_emails = os.environ.get("ADMIN_EMAILS", "").split(",")
    admin_emails = [email.strip().lower() for email in admin_emails if email.strip()]
    
    if member.email in admin_emails:
        flash("Cannot delete other admin users.", "error")
        return redirect(url_for("admin"))
    
    try:
        # Remove member from all projects
        for project in member.projects:
            project.team_members.remove(member)
        
        # Delete the member
        db.session.delete(member)
        safe_commit()
        flash(f"Team member {member.name} ({member.email}) has been deleted.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting member: {str(e)}", "error")
    
    return redirect(url_for("admin"))


@app.route("/update_task/<int:task_id>", methods=["POST"])
def update_task(task_id):
    task = Task.query.get_or_404(task_id)
    try:
        task.progress = float(request.form.get("progress", task.progress))
        if request.form.get("assigned_to_id"):
            task.assigned_to_id = int(request.form.get("assigned_to_id"))
        if request.form.get("planned_cost"):
            task.planned_cost = float(request.form.get("planned_cost"))
        if request.form.get("actual_cost"):
            task.actual_cost = float(request.form.get("actual_cost"))
        safe_commit()
        flash("Task progress updated.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Could not update task: {e}", "error")
    return redirect(url_for("index"))

@app.route("/edit_task/<int:task_id>", methods=["GET", "POST"], strict_slashes=False)
@app.route("/edit_task/<int:task_id>/", methods=["GET", "POST"], strict_slashes=False)
def edit_task(task_id):
    task = Task.query.get_or_404(task_id)
    project = Project.query.get(task.project_id)
    if not project:
        flash("Project not found", "error")
        return redirect(url_for("index"))

    if request.method == "POST":
        try:
            task.name = request.form.get("name", task.name)
            task.progress = float(request.form.get("progress", task.progress))
            task.activity_description = request.form.get("activity_description", task.activity_description)
            task.dependencies = request.form.get("dependencies", task.dependencies)
            task.task_category = request.form.get("task_category", task.task_category)

            start_date_raw = request.form.get("start_date")
            end_date_raw = request.form.get("end_date")
            duration_raw = request.form.get("duration_days") or request.form.get("duration")

            start_date_obj = parse_date(start_date_raw) if start_date_raw else task.start_date
            end_date_obj = parse_date(end_date_raw) if end_date_raw else task.end_date
            duration_days = int(duration_raw) if duration_raw not in (None, "") else task.duration_days

            if start_date_obj and duration_days:
                end_date_obj = calculate_end_date(start_date_obj, duration_days)
            if start_date_obj and end_date_obj and duration_days is None:
                duration_days = max(1, (end_date_obj - start_date_obj).days)

            project_duration = get_project_duration_days(project)
            if project_duration and duration_days and duration_days > project_duration:
                flash("Task duration exceeds project duration!", "danger")
                return redirect(url_for("edit_task", task_id=task.id))

            if project.start_date and project.end_date and start_date_obj and end_date_obj:
                if start_date_obj < project.start_date or end_date_obj > project.end_date:
                    flash("Task dates must stay within the project date range.", "danger")
                    return redirect(url_for("edit_task", task_id=task.id))

            task.start_date = start_date_obj
            task.end_date = end_date_obj
            task.duration_days = duration_days

            if request.form.get("assigned_to_id"):
                task.assigned_to_id = int(request.form.get("assigned_to_id"))
            if request.form.get("planned_cost"):
                task.planned_cost = float(request.form.get("planned_cost"))
            if request.form.get("actual_cost"):
                task.actual_cost = float(request.form.get("actual_cost"))

            safe_commit()

            for image in request.files.getlist("images"):
                if image and image.filename and allowed_file(image.filename, "image"):
                    filename = secure_filename(f"{datetime.now().timestamp()}_{image.filename}")
                    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                    image.save(filepath)
                    db.session.add(MediaFile(filename=filename, filepath=filepath, file_type="image", task_id=task.id))

            for video in request.files.getlist("videos"):
                if video and video.filename and allowed_file(video.filename, "video"):
                    filename = secure_filename(f"{datetime.now().timestamp()}_{video.filename}")
                    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                    video.save(filepath)
                    db.session.add(MediaFile(filename=filename, filepath=filepath, file_type="video", task_id=task.id))

            safe_commit()
            flash("Task updated successfully!", "success")
            return redirect(url_for("edit_task", task_id=task.id))
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Edit task error: {e}")
            flash(f"Error updating task: {str(e)}", "error")
            return redirect(url_for("edit_task", task_id=task.id))

    media_files = MediaFile.query.filter_by(task_id=task.id).all()
    return render_template("edit_task.html", task=task, project=project, media_files=media_files)

@app.route("/gantt_data")
def gantt_data():
    tasks = Task.query.all()
    data = []
    for t in tasks:
        if t.start_date and t.end_date:
            project = Project.query.get(t.project_id)
            data.append({
                "id": t.id,
                "name": t.name,
                "start": t.start_date.strftime("%Y-%m-%d"),
                "end": t.end_date.strftime("%Y-%m-%d"),
                "progress": t.progress,
                "project_id": t.project_id,
                "project_name": project.name if project else f"Project {t.project_id}",
            })
    return jsonify(data)

@app.route("/update_task_gantt/<int:task_id>", methods=["POST"])
def update_task_gantt(task_id):
    task = Task.query.get_or_404(task_id)
    data = request.get_json(force=True)
    try:
        if "start_date" in data and "end_date" in data:
            start = datetime.strptime(data["start_date"], "%Y-%m-%d").date()
            end = datetime.strptime(data["end_date"], "%Y-%m-%d").date()
            if end < start:
                return jsonify({"status": "error", "message": "End date cannot be before start date"}), 400
            task.start_date = start
            task.end_date = end
            task.duration_days = max(1, (end - start).days)

        if "progress" in data:
            task.progress = float(data["progress"])

        safe_commit()
        return jsonify({"status": "ok", "message": "Task updated successfully"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 400

@app.route("/s_curve_data/<int:project_id>")
def s_curve_data(project_id):
    tasks = Task.query.filter_by(project_id=project_id).all()
    rows = []
    for t in tasks:
        if t.start_date:
            planned, actual = task_cost(t)
            rows.append({"date": t.start_date.strftime("%Y-%m-%d"), "planned": planned, "actual": actual})

    rows.sort(key=lambda x: x["date"])
    c_planned = 0
    c_actual = 0
    result = []
    for row in rows:
        c_planned += row["planned"]
        c_actual += row["actual"]
        result.append({"date": row["date"], "planned": round(c_planned, 2), "actual": round(c_actual, 2)})
    return jsonify(result)

@app.route("/team_performance/<int:project_id>")
def team_performance(project_id):
    project = Project.query.get_or_404(project_id)
    tasks = Task.query.filter_by(project_id=project_id).all()
    result = []
    for member in project.team_members:
        assigned = [t for t in tasks if t.assigned_to_id == member.id]
        total = len(assigned)
        completed = len([t for t in assigned if float(t.progress or 0) >= 100])
        delayed = len([t for t in assigned if t.end_date and project.end_date and t.end_date > project.end_date])
        avg_progress = round(sum(float(t.progress or 0) for t in assigned) / total, 1) if total else 0
        result.append({
            "name": member.name,
            "email": member.email,
            "completed": completed,
            "delayed": delayed,
            "total": total,
            "avg_progress": avg_progress,
        })
    return jsonify(result)

# ================= DELETE ROUTES =================
@app.route("/delete_media/<int:media_id>")
def delete_media(media_id):
    media = MediaFile.query.get_or_404(media_id)
    task_id = media.task_id
    try:
        if media.filepath and os.path.exists(media.filepath):
            os.remove(media.filepath)
        db.session.delete(media)
        safe_commit()
        flash("Media file deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Could not delete media: {str(e)}", "error")
    return redirect(url_for("edit_task", task_id=task_id))

@app.route("/delete_task/<int:task_id>")
def delete_task(task_id):
    task = Task.query.get_or_404(task_id)
    try:
        for media in MediaFile.query.filter_by(task_id=task_id).all():
            if media.filepath and os.path.exists(media.filepath):
                os.remove(media.filepath)
            db.session.delete(media)
        db.session.delete(task)
        safe_commit()
        flash("Task deleted.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Could not delete task: {str(e)}", "error")
    return redirect(url_for("index"))

@app.route("/delete_all_tasks/<int:project_id>")
def delete_all_tasks(project_id):
    Project.query.get_or_404(project_id)
    try:
        tasks = Task.query.filter_by(project_id=project_id).all()
        for task in tasks:
            for media in MediaFile.query.filter_by(task_id=task.id).all():
                if media.filepath and os.path.exists(media.filepath):
                    os.remove(media.filepath)
                db.session.delete(media)
            db.session.delete(task)
        safe_commit()
        flash("All tasks deleted.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Could not delete tasks: {str(e)}", "error")
    return redirect(url_for("index"))

@app.route("/delete_project/<int:project_id>")
def delete_project(project_id):
    project = Project.query.get_or_404(project_id)
    try:
        tasks = Task.query.filter_by(project_id=project.id).all()
        for task in tasks:
            for media in MediaFile.query.filter_by(task_id=task.id).all():
                if media.filepath and os.path.exists(media.filepath):
                    os.remove(media.filepath)
                db.session.delete(media)
        Task.query.filter_by(project_id=project.id).delete()
        db.session.delete(project)
        safe_commit()
        flash("Project deleted.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Could not delete project: {str(e)}", "error")
    return redirect(url_for("index"))

# ================= REPORT ROUTES =================
@app.route("/project_report/<int:project_id>")
def project_report(project_id):
    project = Project.query.get_or_404(project_id)
    tasks = Task.query.filter_by(project_id=project_id).all()
    doc = BytesIO()
    doc.write(f"{project.name} Report\n".encode())
    for t in tasks:
        doc.write(f"{t.name}: {t.progress}%\n".encode())
    doc.seek(0)
    return send_file(doc, as_attachment=True, download_name=f"{project.name}_report.txt")

@app.route("/generate_report/<int:project_id>")
def generate_report(project_id):
    from docx import Document
    from docx.shared import Inches
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    project = Project.query.get_or_404(project_id)
    tasks = Task.query.filter_by(project_id=project_id).all()

    doc = Document()
    doc.add_heading(f"{project.name} - Progress Report", 0)
    doc.add_paragraph(f"Type: {project.project_type}")
    if project.start_date:
        doc.add_paragraph(f"Start Date: {project.start_date}")
    if project.end_date:
        doc.add_paragraph(f"End Date: {project.end_date}")
    doc.add_paragraph(f"Overall Completion: {project.completion:.1f}%")

    doc.add_heading("Task Progress", level=1)
    for t in tasks:
        doc.add_paragraph(f"{t.name}: {t.progress}%")
        media_files = MediaFile.query.filter_by(task_id=t.id).all()
        if media_files:
            doc.add_paragraph(f"  Attached files: {len(media_files)}")
            for media in media_files:
                doc.add_paragraph(f"    - {media.filename} ({media.file_type})")

    try:
        plt.figure(figsize=(6, 4))
        plt.bar([t.name for t in tasks], [float(t.progress or 0) for t in tasks])
        plt.ylabel("Progress (%)")
        plt.xticks(rotation=45, ha="right")
        plt.tight_layout()
        img_stream1 = BytesIO()
        plt.savefig(img_stream1, format="png")
        img_stream1.seek(0)
        doc.add_picture(img_stream1, width=Inches(5))
        plt.close()

        plt.figure(figsize=(4, 4))
        completion = max(0.0, min(100.0, project.completion or 0))
        plt.pie([completion, max(0, 100 - completion)], labels=["Completed", "Remaining"], autopct="%1.1f%%")
        plt.title("Overall Completion")
        img_stream2 = BytesIO()
        plt.savefig(img_stream2, format="png")
        img_stream2.seek(0)
        doc.add_picture(img_stream2, width=Inches(4))
        plt.close()
    except Exception as e:
        print(f"Error generating charts: {e}")

    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return send_file(
        word_stream,
        as_attachment=True,
        download_name=f"{project.name}_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

@app.route("/invite_settings")
def invite_settings():
    email_configured = bool(app.config["MAIL_USERNAME"] and app.config["MAIL_PASSWORD"])
    return render_template(
        "invite_settings.html",
        email_configured=email_configured,
        mail_server=app.config["MAIL_SERVER"],
        mail_port=app.config["MAIL_PORT"],
        mail_username=app.config["MAIL_USERNAME"],
    )

@app.route("/project_members/<int:project_id>/resend/<int:member_id>")
def resend_invitation(project_id, member_id):
    project = Project.query.get_or_404(project_id)
    member = TeamMember.query.get_or_404(member_id)

    if member not in project.team_members:
        flash("Member is not part of this project", "error")
        return redirect(url_for("project_members", project_id=project.id))

    token_data = f"{member.id}_{project.id}_{datetime.utcnow().timestamp()}_{secrets.token_hex(8)}"
    invite_token = hashlib.sha256(token_data.encode()).hexdigest()
    member.invite_token = invite_token
    member.token_expiry = datetime.utcnow() + timedelta(days=7)
    safe_commit()

    email_sent = send_invitation_email(member.email, project.name, project.id, invite_token)
    flash("Invitation resent to member." if email_sent else "Failed to resend invitation.", "success" if email_sent else "error")
    return redirect(url_for("project_members", project_id=project.id))

@app.route("/api/project/<int:project_id>/members")
def get_project_members_api(project_id):
    project = Project.query.get_or_404(project_id)
    members = []
    for member in project.team_members:
        members.append({
            "id": member.id,
            "email": member.email,
            "name": member.name,
            "joined_at": member.joined_at.strftime("%Y-%m-%d %H:%M:%S") if getattr(member, "joined_at", None) else None
        })
    return jsonify({
        "success": True,
        "project_id": project.id,
        "project_name": project.name,
        "members": members,
        "total_members": len(members)
    })

# ================= OTHER ROUTES =================
@app.route("/drone_view")
def drone_view():
    try:
        return render_template("drone_view.html")
    except Exception:
        return """
        <!DOCTYPE html>
        <html>
        <head><title>Drone View</title></head>
        <body>
            <h1>Drone View Coming Soon</h1>
            <a href="/">Back to Dashboard</a>
        </body>
        </html>
        """

@app.route("/static/uploads/<path:filename>")
def uploaded_file(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)

# ================= ERROR HANDLERS =================
@app.errorhandler(404)
def not_found(e):
    return "<h1>404</h1><p>Page not found</p><a href='/'>Home</a>", 404

@app.errorhandler(500)
def internal_error(e):
    db.session.rollback()
    return "<h1>500</h1><p>Internal server error</p><a href='/'>Home</a>", 500

@app.teardown_appcontext
def shutdown_session(exception=None):
    db.session.remove()
    gc.collect()

# ================= STARTUP =================
with app.app_context():
    try:
        db.create_all()
    except Exception as e:
        print(f"Error creating tables: {e}")

    if scheduler and not scheduler.running:
        try:
            scheduler.add_job(send_weekly_reports, "cron", day_of_week="mon", hour=8, minute=0)
            scheduler.start()
            print("Weekly report scheduler started")
        except Exception as e:
            print(f"Scheduler start failed: {e}")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=False)